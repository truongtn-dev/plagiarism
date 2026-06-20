from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from checker.models import ParagraphResult, RiskLevel, ScanReport

COPYRIGHT = "Phần mềm được viết bởi Nguyễn Thành Trương"
APP_NAME = "IEEE Plagiarism Checker"

RISK_LABELS = {
    RiskLevel.LOW: "An toàn",
    RiskLevel.MEDIUM: "Theo dõi",
    RiskLevel.HIGH: "Cao",
    RiskLevel.CRITICAL: "Nghiêm trọng",
}

RISK_COLORS = {
    RiskLevel.LOW: RGBColor(0x22, 0xC5, 0x5E),
    RiskLevel.MEDIUM: RGBColor(0xF5, 0x9E, 0x0B),
    RiskLevel.HIGH: RGBColor(0xEF, 0x44, 0x44),
    RiskLevel.CRITICAL: RGBColor(0xDC, 0x26, 0x26),
}


def _set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def _add_label_value(doc: Document, label: str, value: str, value_color=None):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    _set_run_font(r1, bold=True)
    r2 = p.add_run(value)
    _set_run_font(r2, color=value_color)
    return p


def export_report_docx(report: ScanReport, min_risk: str = "medium+") -> bytes:
    """Generate Word report with flagged paragraphs and fix suggestions."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BÁO CÁO KIỂM TRA ĐẠO VĂN\n")
    _set_run_font(r, size=18, bold=True, color=RGBColor(0x1E, 0x40, 0xAF))
    r2 = title.add_run(f"{APP_NAME}\n")
    _set_run_font(r2, size=12, color=RGBColor(0x64, 0x74, 0x8B))
    r3 = title.add_run(COPYRIGHT)
    _set_run_font(r3, size=10, italic=True, color=RGBColor(0x94, 0xA3, 0xB8))

    doc.add_paragraph()

    # Summary box
    _add_heading(doc, "Tóm tắt kết quả", level=2)
    _add_label_value(doc, "File gốc", report.filename)
    _add_label_value(doc, "Thời gian quét", datetime.now().strftime("%d/%m/%Y %H:%M"))
    _add_label_value(doc, "Chế độ quét", report.scan_mode)
    _add_label_value(doc, "Thời lượng", f"{report.duration_seconds}s")
    _add_label_value(
        doc,
        "Web match (strict)",
        f"{report.plagiarism_percent}%",
        RGBColor(0xDC, 0x26, 0x26),
    )
    _add_label_value(
        doc,
        "IEEE risk estimate",
        f"{report.ieee_estimate_percent}%",
        RGBColor(0xDC, 0x26, 0x26) if report.ieee_estimate_percent >= 30 else RGBColor(0xF5, 0x9E, 0x0B),
    )
    _add_label_value(
        doc,
        "Độ nguyên bản (web)",
        f"{report.originality_percent}%",
        RGBColor(0x22, 0xC5, 0x5E),
    )
    _add_label_value(doc, "Đoạn đã phân tích", str(report.analyzed_paragraphs))

    p = doc.add_paragraph()
    r = p.add_run("Phân bố rủi ro: ")
    _set_run_font(r, bold=True)
    parts = []
    for level in RiskLevel:
        count = report.risk_summary.get(level.value, 0)
        if count:
            parts.append(f"{RISK_LABELS[level]}: {count}")
    p.add_run(" · ".join(parts))

    doc.add_page_break()

    # Filter paragraphs
    risk_order = {
        "high+": [RiskLevel.HIGH, RiskLevel.CRITICAL],
        "medium+": [RiskLevel.MEDIUM, RiskLevel.HIGH, RiskLevel.CRITICAL],
        "all": list(RiskLevel),
    }
    allowed = risk_order.get(min_risk, list(RiskLevel))

    flagged = [
        p
        for p in report.paragraphs
        if p.risk in allowed and (p.similarity > 0 or p.risk != RiskLevel.LOW)
    ]
    flagged.sort(key=lambda p: (-p.similarity, p.index))

    _add_heading(doc, f"Chi tiết cần sửa ({len(flagged)} đoạn)", level=2)

    intro = doc.add_paragraph()
    intro.add_run(
        "Mở file Word này song song với bài gốc để xem từng lỗi và gợi ý sửa. "
        "Các đoạn được sắp xếp theo mức độ trùng khớp giảm dần."
    )

    if not flagged:
        doc.add_paragraph("Không phát hiện đoạn văn có rủi ro đáng kể ở mức lọc hiện tại.")
    else:
        for i, para in enumerate(flagged, 1):
            _add_paragraph_report(doc, para, i)

    # Footer copyright on last section
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run(f"\n— {COPYRIGHT} —")
    _set_run_font(rf, size=9, italic=True, color=RGBColor(0x94, 0xA3, 0xB8))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_paragraph_report(doc: Document, para: ParagraphResult, order: int):
    """Add one flagged paragraph section to the Word report."""
    doc.add_paragraph()

    # Section header
    h = doc.add_paragraph()
    hr = h.add_run(
        f"#{order} · Đoạn {para.index} · {para.section} · "
        f"{para.similarity}% trùng khớp · {RISK_LABELS[para.risk]}"
    )
    _set_run_font(hr, size=12, bold=True, color=RISK_COLORS.get(para.risk))

    # Original text
    p = doc.add_paragraph()
    lbl = p.add_run("Nội dung gốc:\n")
    _set_run_font(lbl, bold=True, size=10)
    body = p.add_run(para.text)
    _set_run_font(body, size=10)
    p.paragraph_format.left_indent = Inches(0.2)

    # Sources
    if para.sources:
        p = doc.add_paragraph()
        p.add_run("Nguồn trùng khớp:\n").bold = True
        for src in para.sources[:5]:
            sp = doc.add_paragraph(style="List Bullet")
            sp.add_run(f"{src.title} ({src.similarity}%)")
            if src.url:
                sp.add_run(f"\n{src.url}")
            if src.matched_text:
                sp.add_run(f"\nĐoạn khớp: «{src.matched_text[:200]}»")

    # Internal duplicates
    if para.internal_duplicates:
        p = doc.add_paragraph()
        p.add_run("Trùng lặp nội bộ:\n").bold = True
        for dup in para.internal_duplicates:
            dp = doc.add_paragraph(style="List Bullet")
            dp.add_run(dup[:300])

    # Suggestions — highlighted for editing
    if para.suggestions:
        p = doc.add_paragraph()
        sg = p.add_run("Gợi ý sửa đổi:\n")
        _set_run_font(sg, bold=True, color=RGBColor(0x1D, 0x4E, 0xD8))
        for s in para.suggestions:
            sp = doc.add_paragraph(style="List Bullet")
            sr = sp.add_run(s)
            _set_run_font(sr, size=10, color=RGBColor(0x1E, 0x3A, 0x8A))

    # Separator
    sep = doc.add_paragraph()
    sep.add_run("─" * 60)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
