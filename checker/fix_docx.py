from __future__ import annotations

import io
import re
import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from checker.citations import (
    CITATION_LEAK_RE,
    apply_paraphrase_in_place,
    extract_parts,
    parts_with_placeholders,
)
from checker.export_docx import APP_NAME, COPYRIGHT
from checker.models import ParagraphResult, RiskLevel, ScanReport
from checker.paraphraser import paraphrase_for_similarity

PLACEHOLDER_RE = re.compile(r"\uE000\d+\uE001")

SKIP_SECTIONS = {
    "Index Terms",
    "References",
    "Header",
    "Figures",
    "Tables",
}


@dataclass
class FixStats:
    paragraphs_modified: int = 0
    paragraphs_unchanged: int = 0
    paragraphs_skipped: int = 0
    paragraphs_failed: int = 0
    modified_indices: list[int] = field(default_factory=list)


def _should_rewrite(para: ParagraphResult, min_similarity: float) -> bool:
    if para.risk not in (RiskLevel.MEDIUM, RiskLevel.HIGH, RiskLevel.CRITICAL):
        return False
    if para.similarity < min_similarity:
        return False
    if para.section in SKIP_SECTIONS:
        return False
    skip_hints = ("Bỏ qua", "Metadata", "References", "email", "tác giả")
    if para.suggestions and any(h in para.suggestions[0] for h in skip_hints):
        return False
    text = para.text.strip()
    if re.match(r"^H[1-4]:", text):
        return False
    if re.match(r"^(TABLE|Table)\s", text):
        return False
    if re.match(r"^[IVXLC]+\.", text) and len(text) < 80:
        return False
    if re.match(r"^\[\d+\]", text):
        return False
    return True


def _backup_runs(paragraph) -> list[tuple[str, object]]:
    return [(r.text, r.font.color.rgb if r.font.color else None) for r in paragraph.runs]


def _restore_runs(paragraph, backup: list[tuple[str, object]]):
    for run, (text, color) in zip(paragraph.runs, backup):
        run.text = text
        if color is not None:
            run.font.color.rgb = color


def _rewrite_paragraph_text(
    paragraph,
    para_result: ParagraphResult,
    min_similarity: float,
) -> bool:
    if not _should_rewrite(para_result, min_similarity):
        return False

    parts = extract_parts(paragraph)
    if not parts:
        return False

    plain_with_ph, _fields = parts_with_placeholders(parts)
    if not plain_with_ph.strip():
        return False

    matched = para_result.sources[0].matched_text if para_result.sources else ""
    new_plain = paraphrase_for_similarity(
        plain_with_ph,
        para_result.similarity,
        matched or "",
    )

    orig_tokens = PLACEHOLDER_RE.findall(plain_with_ph)
    new_tokens = PLACEHOLDER_RE.findall(new_plain)
    if orig_tokens != new_tokens:
        return False

    orig_clean = PLACEHOLDER_RE.sub("", plain_with_ph).strip()
    new_clean = PLACEHOLDER_RE.sub("", new_plain).strip()
    if orig_clean == new_clean:
        return False

    if CITATION_LEAK_RE.search(new_clean):
        return False

    backup = _backup_runs(paragraph)
    ok = apply_paraphrase_in_place(
        paragraph,
        new_plain,
        plain_with_ph,
        highlight_changes=True,
    )
    if not ok or CITATION_LEAK_RE.search(paragraph.text):
        _restore_runs(paragraph, backup)
        return False

    return True


def generate_fixed_docx(
    source_path: str,
    report: ScanReport,
    min_similarity: float = 45.0,
) -> tuple[bytes, FixStats]:
    """
    Copy source docx, paraphrase flagged paragraphs in-place.
    Mendeley citation field runs are never modified.
    """
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    shutil.copy2(source_path, tmp.name)

    doc = Document(tmp.name)
    Path(tmp.name).unlink(missing_ok=True)

    result_map = {p.index: p for p in report.paragraphs}
    stats = FixStats()

    for doc_idx, paragraph in enumerate(doc.paragraphs):
        para_num = doc_idx + 1
        if para_num not in result_map:
            continue
        para_result = result_map[para_num]
        if not paragraph.text.strip():
            stats.paragraphs_skipped += 1
            continue

        try:
            if _rewrite_paragraph_text(paragraph, para_result, min_similarity):
                stats.paragraphs_modified += 1
                stats.modified_indices.append(para_num)
            else:
                stats.paragraphs_unchanged += 1
        except Exception:
            stats.paragraphs_failed += 1

    doc.add_paragraph()
    note = doc.add_paragraph()
    r1 = note.add_run(
        f"[{APP_NAME}] Phần màu đỏ đã được paraphrase. "
        f"Citation Mendeley giữ nguyên — mở Word > Mendeley > Refresh nếu cần. {COPYRIGHT}"
    )
    from docx.shared import Pt, RGBColor

    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r1.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), stats
